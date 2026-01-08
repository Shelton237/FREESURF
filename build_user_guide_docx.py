from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

OUTPUT_PATH = r'docs/guide_utilisateur_freesurf.docx'

document = Document()
document.core_properties.title = 'Guide utilisateur FREESURF / CuWiP'
document.core_properties.subject = 'Mode opératoire Backoffice, Technicien et Portail client'
document.core_properties.language = 'fr-FR'

title = document.add_paragraph()
title_run = title.add_run('Guide utilisateur\nFREESURF / CuWiP')
title_run.bold = True
title_run.font.size = Pt(24)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

subtitle = document.add_paragraph('Version initiale — Décembre 2025')
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

document.add_page_break()

sections = [
    {
        'heading': '1. Mise en route',
        'body': [
            'Ce projet repose sur Laravel 12, Inertia/Vue 3 et une Progressive Web App. Avant toute utilisation, vérifiez que l’environnement de développement ou de production est correctement initialisé.'
        ],
        'bullets': [
            'Prérequis: PHP 8.2+, Composer, Node 20+, Redis pour Horizon, Docker (optionnel pour la prod).',
            'Installation locale: cd backend → composer install → cp .env.example .env → php artisan key:generate → npm ci → npm run dev → php artisan serve.',
            'Tâches de fond: lancer Horizon avec «php artisan horizon» et laisser la tâche planifiée «invoices:generate-monthly» créer les factures chaque 1er à 01h.',
            'Déploiement: utiliser le workflow GitHub Actions et/ou docker compose (services php-fpm, nginx, horizon, scheduler, redis, db, Traefik).',
        ],
    },
    {
        'heading': '2. Backoffice (Managers)',
        'body': [
            'Accessible via «/backoffice» pour les utilisateurs authentifiés avec e-mail vérifié et rôle backoffice. Cette zone couvre toute la gestion référentielle et opérationnelle.'
        ],
        'bullets': [
            'BTS: créer des stations (code généré automatiquement), suivre la charge et les clients associés.',
            'Clients: enregistrer les prospects (coordonnées, GPS, photos, partenaire), saisir les études d’éligibilité et marquer les installations terminées.',
            'Ordres de travail: créer/filtrer des interventions survey/install/maintenance, affecter un technicien et consulter les synthèses de surveys.',
            'Demandes portail: suivre les demandes d’abonnement/réabonnement soumises depuis l’espace client.',
            'SAV: ouvrir ou modifier des tickets (incident/assistance/réclamation), assigner un intervenant, planifier un suivi.',
            'Administration: créer/éditer les comptes utilisateurs (roles backoffice ou technicien) et réinitialiser les mots de passe.',
        ],
    },
    {
        'heading': '3. Espace technicien',
        'body': [
            'L’espace «/tech» est réservé aux comptes rôle technicien; il regroupe le tableau de bord terrain, la création d’études et le suivi des tickets SAV.'
        ],
        'bullets': [
            'Tableau de bord: visualiser les interventions assignées, filtrer par statut (pending = assigned/accepted, on_site, completed) et accéder aux fiches.',
            'Nouvelle étude: rechercher un client existant, planifier une visite survey et générer un WorkOrder directement depuis le terrain.',
            'Fiche intervention: démarrer la visite (saisie manuelle ou GPS auto), clôturer avec date/commentaire, renseigner le résultat survey ou l’installation, téléverser des photos.',
            'Automatisme: la clôture «install» passe le client à actif; une étude «survey» met à jour l’éligibilité et peut créer un ticket SAV de suivi.',
            'SAV terrain: consulter/filtrer les tickets non assignés ou en cours, s’auto-assigner, mettre à jour le statut, ajouter des notes et planifier un rappel.',
        ],
    },
    {
        'heading': '4. Portail client',
        'body': [
            'Les clients disposent d’un site portail (routes /portal et /portal/compte) pour initier leurs démarches et suivre leur dossier.'
        ],
        'bullets': [
            'Demandes d’abonnement/réabonnement: formulaire public demandant identité, coordonnées GPS, pièces et création/mise à jour du compte (mot de passe requis). Contrôle du code client et des impayés pour un réabonnement.',
            'Authentification: connexion par téléphone/mot de passe, session stockée côté portail avec bouton de déconnexion.',
            'Dashboard: vue synthétique des demandes soumises, des clients liés (statut + factures impayées), des dernières factures et des tickets SAV en cours.',
            'Lier un client existant: saisir code + téléphone pour rattacher un dossier et rapatrier ses factures dans l’espace client.',
            'Factures: filtrer par statut, ouvrir le détail, télécharger le PDF (stocké ou généré via Dompdf), consulter l’historique de paiements.',
            'Timeline des demandes: suivre les étapes (soumise → en étude → planification → installée → active) et annuler tant que la demande n’est pas installée.',
            'SAV self-service: soumettre un ticket (incident/assistance/réclamation) sur un client lié; le canal est noté «portal» et le backoffice prend le relais.',
        ],
    },
    {
        'heading': '5. Automatisation & support',
        'body': [
            'Plusieurs briques soutiennent la production continue, la supervision et l’expérience mobile.'
        ],
        'bullets': [
            'Facturation mensuelle: surveiller la commande «invoices:generate-monthly» et relancer manuellement si besoin après avoir appliqué les exclusions.',
            'Queues & Horizon: Redis traite les jobs (notifications, PDF, etc.); Horizon fournit un monitoring temps réel.',
            'PWA: l’interface Inertia/Vue est installable sur mobile, désactivable en CI avec la variable SKIP_PWA.',
            'Stack Docker/Traefik: fournie pour la prod (nginx, PHP-FPM, Redis, Horizon, scheduler, DB, Traefik + certificats Let’s Encrypt).',
        ],
    },
]

for section in sections:
    document.add_heading(section['heading'], level=1)
    for paragraph in section.get('body', []):
        document.add_paragraph(paragraph)
    for bullet in section.get('bullets', []):
        document.add_paragraph(bullet, style='List Bullet')
    document.add_paragraph('')

os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
document.save(OUTPUT_PATH)
print('Wrote', OUTPUT_PATH)
