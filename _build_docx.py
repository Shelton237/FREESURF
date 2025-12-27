from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re, os

md_path = r'docs/cahier_de_specifications.md'
docx_path = r'docs/cahier_de_specifications.docx'

document = Document()
document.core_properties.title = 'Cahier de Spécification des Besoins – CuWiP'
document.core_properties.subject = 'Spécifications fonctionnelles et non-fonctionnelles'
document.core_properties.language = 'fr-FR'

# Page de garde
p = document.add_paragraph()
r = p.add_run('Cahier de Spécification des Besoins\nCuWiP (Customer Wireless Provider)')
r.bold = True
r.font.size = Pt(20)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

for img in ['logo.jpeg', 'logo_black.jpeg']:
    if os.path.exists(img):
        document.add_picture(img, width=Inches(2.5))
        document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        break

document.add_page_break()

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.read().splitlines()

for line in lines:
    raw = line.rstrip('\n')
    if not raw.strip():
        document.add_paragraph('')
        continue
    if raw.startswith('# '):
        document.add_heading(raw[2:].strip(), level=1)
        continue
    if raw.startswith('## '):
        document.add_heading(raw[3:].strip(), level=2)
        continue
    if raw.startswith('### '):
        document.add_heading(raw[4:].strip(), level=3)
        continue
    m = re.match(r'^-\s+(.*)', raw)
    if m:
        document.add_paragraph(m.group(1), style='List Bullet')
        continue
    m = re.match(r'^(\d+)\.\s+(.*)', raw)
    if m:
        document.add_paragraph(m.group(2), style='List Number')
        continue
    if set(raw.strip()) == {'-'} and len(raw.strip()) >= 3:
        document.add_paragraph('')
        continue
    document.add_paragraph(raw)

os.makedirs(os.path.dirname(docx_path), exist_ok=True)
document.save(docx_path)
print('Wrote', docx_path)
